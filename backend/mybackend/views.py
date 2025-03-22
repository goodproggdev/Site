from rest_framework.decorators import api_view, parser_classes
from rest_framework.parsers import MultiPartParser
from rest_framework.response import Response
from django.http import JsonResponse
from django.views.decorators.csrf import ensure_csrf_cookie, csrf_protect
from django.core.mail import send_mail
from django.views.decorators.csrf import csrf_exempt
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile
import os
import docx
import PyPDF2
import socket
import smtplib
import json
import io
from django.core.files.base import ContentFile

@api_view(['POST'])
@parser_classes([MultiPartParser])
def analyze_cv(request):
	if 'document' not in request.data:
		return Response({'error':'Nessun file inviato'}, status=400)
	document = request.data['document']
	# Esegui l'analisi del documento qui
	result = {'message':'Documento analizzato con successo'}
	return Response(result)

@csrf_exempt
def upload_file(request):
    if request.method == 'POST' and request.FILES.get('file'):
        file_obj = request.FILES['file']

        print("File:", file_obj)

        try:
            original_file_path = default_storage.save('uploads/original/' + file_obj.name, ContentFile(file_obj.read()))
            original_file_url = default_storage.url(original_file_path)

            print(f"File salvato in: {original_file_path}") # Debug

            if default_storage.exists(original_file_path): # Debug
                print("File salvato con successo")
            else:
                print("Errore: il file non è stato salvato")

            text = extract_text(original_file_path, file_obj.name)

            txt_file_path = default_storage.save('uploads/text/' + os.path.splitext(file_obj.name)[0] + '.txt', ContentFile(text.encode('utf-8')))
            txt_file_url = default_storage.url(txt_file_path)

            return JsonResponse({
                'message': 'File caricato e testo estratto con successo.',
                'original_file_url': original_file_url,
                'txt_file_url': txt_file_url
            })
        except Exception as e:
            print(f"Errore durante il salvataggio del file: {e}")
            return JsonResponse({'message': f'Errore durante il salvataggio del file: {e}'}, status=500)
    return JsonResponse({'message': 'Errore nel caricamento del file.'}, status=400)

def extract_text(file_path, file_name):
    try:
        with default_storage.open(file_path, 'rb') as file: #usa default_storage.open
            if file_name.endswith('.pdf'):
                reader = PyPDF2.PdfReader(io.BytesIO(file.read())) #legge il file come bytes
                text = ''
                for page in reader.pages:
                    text += page.extract_text() or ''
                return text
            elif file_name.endswith('.docx'):
                doc = docx.Document(io.BytesIO(file.read())) #legge il file come bytes
                text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                return text
            elif file_name.endswith('.txt'):
                with default_storage.open(file_path, 'r', encoding='utf-8') as file: #usa default_storage.open
                    return file.read()
            else:
                return "File type not supported."
    except (PyPDF2.errors.PdfReadError, docx.opc.exceptions.PackageNotFoundError) as e:
        return f"Error processing file: {e}"
    except FileNotFoundError:
        return "File not found."
    except Exception as e:
        return f"An unexpected error occurred: {e}"
    try:
        if file_name.endswith('.pdf'):
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text = ''
                for page in reader.pages:
                    text += page.extract_text() or ''
            return text
        elif file_name.endswith('.docx'):
            doc = docx.Document(file_path)
            text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            return text
        elif file_name.endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        else:
            return "File type not supported."
    except (PyPDF2.errors.PdfReadError, docx.opc.exceptions.PackageNotFoundError) as e:
        return f"Error processing file: {e}"
    except FileNotFoundError:
        return "File not found."
    except Exception as e:
        return f"An unexpected error occurred: {e}"
    
@ensure_csrf_cookie
@csrf_protect
def contact_view(request):
    def check_internet():
        try:
            socket.create_connection(("8.8.8.8", 53), timeout=5)
            return True
        except OSError:
            return False

    if request.method == 'POST':
        try:
            if not check_internet():
                return JsonResponse({'error': 'Nessuna connessione internet'}, status=503)

            data = json.loads(request.body)
            
            email = data.get('email', '').strip()
            subject = data.get('subject', '').strip()
            message = data.get('message', '').strip()

            if not all([email, subject, message]):
                return JsonResponse({'error': 'Compila tutti i campi'}, status=400)

            try:
                send_mail(
                    subject=f"{subject} - From {email}",
                    message=message,
                    from_email=email,
                    recipient_list=['sitiegestionali@gmail.com'],
                    fail_silently=False,
                )
            except smtplib.SMTPException as e:
                error_code = e.smtp_code if hasattr(e, 'smtp_code') else '000'
                return JsonResponse({
                    'error': f"Errore SMTP ({error_code}): Impossibile inviare l'email",
                    'details': str(e)
                }, status=503)

            return JsonResponse({'status': 'success'})

        except json.JSONDecodeError:
            return JsonResponse({'error': 'Formato dati non valido'}, status=400)
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)

    return JsonResponse({'status': 'csrf cookie set'})